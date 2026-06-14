"""Cover letter pipeline tests: tailor service, renderer, linter, pipeline.

Covers the P5 additions to the tailor pipeline. Uses MockLLMClient for
LLM interactions and patched service dependencies for DB side effects.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime
from typing import Any
from unittest.mock import MagicMock

import pytest
from docx import Document

from app.models.ats_lint import LintResult, LintViolation
from app.models.experience import (
    OptimizedDoc,
    OptimizedPayload,
    Outcome,
    PreferencesPayload,
    Role,
    Skill,
)
from app.models.tailor import (
    ContactInfo,
    CoverLetterParagraph,
    TailoredCoverLetter,
)
from app.services.ats_lint import lint_docx
from app.services.docx.renderer import render_cover_letter_docx, render_docx
from app.services.llm import cost_log as cost_log_mod
from app.services.llm.mock import MockLLMClient
from app.services.tailor import pipeline as pipeline_module
from app.services.tailor.pipeline import (
    CoverLetterPipelineLintFailure,
    CoverLetterPipelineSuccess,
    run_cover_letter_pipeline,
)
from app.services.tailor.prompts import COVER_LETTER_SYSTEM
from app.services.tailor.tailor import (
    DEFAULT_COVER_LETTER_PURPOSE,
    build_cover_letter_user_message,
    tailor_cover_letter,
    validate_cover_letter_refs,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _optimized_payload() -> OptimizedPayload:
    return OptimizedPayload(
        summary="Senior FE.",
        roles=[
            Role(
                id="fc",
                company="FightCamp",
                title="Senior Frontend Engineer",
                start="2021-11",
                end="2024-04",
                summary="Led the PDP rebuild.",
                skills=["React"],
                outcome_refs=[],
            )
        ],
        skills=[Skill(name="React"), Skill(name="TypeScript")],
        outcomes=[
            Outcome(
                description="Cut mobile load times from 10s to 2s",
                metric="LCP",
                value="2s",
                role_ref="fc",
            )
        ],
    )


def _optimized_doc() -> OptimizedDoc:
    return OptimizedDoc(
        id="opt-1",
        user_id="test-user",
        prose_doc_id=None,
        version=1,
        payload=_optimized_payload(),
        markdown_view=None,
        source="llm",
        created_at=datetime.now(UTC),
    )


def _contact() -> ContactInfo:
    return ContactInfo(name="Daniel Joffe", email="daniel@example.com")


def _valid_letter(
    *,
    outcome_refs: list[str] | None = None,
    role_refs: list[str] | None = None,
    skill_refs: list[str] | None = None,
) -> TailoredCoverLetter:
    return TailoredCoverLetter(
        contact=_contact(),
        recipient_company="Acme Corp",
        recipient_role="Senior FE",
        salutation="Dear Acme Hiring Team,",
        paragraphs=[
            CoverLetterParagraph(text="I am applying for the Senior FE role at Acme."),
            CoverLetterParagraph(text="At FightCamp I cut mobile load times from 10s to 2s."),
            CoverLetterParagraph(text="I would welcome a conversation."),
        ],
        closing="Sincerely,",
        signature="Daniel Joffe",
        source_outcome_refs=outcome_refs
        if outcome_refs is not None
        else ["Cut mobile load times from 10s to 2s"],
        source_role_refs=role_refs if role_refs is not None else ["fc"],
        source_skill_refs=skill_refs if skill_refs is not None else ["React"],
    )


def _valid_letter_json(**kwargs: Any) -> str:
    return _valid_letter(**kwargs).model_dump_json()


def _inserted_record_row(record_id: str = "rec-1") -> dict[str, Any]:
    letter = _valid_letter()
    return {
        "id": record_id,
        "user_id": None,
        "job_posting_id": None,
        "document_type": "cover_letter",
        "resume_type": "generic",
        "jd_snapshot": "JD text",
        "jd_snapshot_hash": "hash",
        "payload": letter.model_dump(mode="json"),
        "storage_path": None,
        "warnings": [],
        "model": "claude-sonnet-4-6",
        "input_tokens": 100,
        "output_tokens": 80,
        "cost_usd": 0.002,
        "latency_ms": 60,
        "created_at": datetime.now(UTC).isoformat(),
    }


def _make_supabase_mock(*, insert_data: list[dict[str, Any]]) -> MagicMock:
    supabase = MagicMock()
    supabase.table.return_value.insert.return_value.execute.return_value.data = (
        insert_data
    )
    supabase.table.return_value.update.return_value.eq.return_value.execute.return_value.data = []
    supabase.storage.from_.return_value.upload.return_value = None
    return supabase


# ---------------------------------------------------------------------------
# System prompt sanity
# ---------------------------------------------------------------------------


def test_cover_letter_system_prompt_mentions_core_rules() -> None:
    assert "TailoredCoverLetter JSON" in COVER_LETTER_SYSTEM
    assert "source_outcome_refs" in COVER_LETTER_SYSTEM
    assert "source_role_refs" in COVER_LETTER_SYSTEM
    assert "source_skill_refs" in COVER_LETTER_SYSTEM
    assert "Do not reference companies, metrics, dates, or facts absent" in COVER_LETTER_SYSTEM
    assert "No em dashes" in COVER_LETTER_SYSTEM


# ---------------------------------------------------------------------------
# build_cover_letter_user_message
# ---------------------------------------------------------------------------


def test_build_cover_letter_user_message_includes_all_sections() -> None:
    msg = build_cover_letter_user_message(
        optimized=_optimized_payload(),
        job_description="We want a senior FE.",
        company_name="Acme",
        contact=_contact(),
        role_title="Senior FE",
        preferences_text='{"rules": ["present tense"]}',
        annotations_text="EMPHASIZE: role roadmap",
        critique="Lead with the PDP rebuild.",
    )
    for tag in (
        "[OptimizedPayload]",
        "[ContactInfo]",
        "[RecipientCompany] Acme",
        "[RoleTitle] Senior FE",
        "[Preferences]",
        "[Annotations]",
        "[Critique]",
        "[JobDescription]",
    ):
        assert tag in msg


def test_build_cover_letter_user_message_omits_role_title_when_absent() -> None:
    msg = build_cover_letter_user_message(
        optimized=_optimized_payload(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
        role_title=None,
        preferences_text=None,
        annotations_text=None,
        critique=None,
    )
    assert "[RoleTitle]" not in msg
    assert "[Preferences]" not in msg
    assert "[Annotations]" not in msg
    assert "[Critique]" not in msg


# ---------------------------------------------------------------------------
# validate_cover_letter_refs
# ---------------------------------------------------------------------------


def test_validate_drops_unknown_role_ref() -> None:
    letter = _valid_letter(role_refs=["fc", "ghost-role"])
    cleaned, warnings = validate_cover_letter_refs(letter, _optimized_payload())
    assert cleaned.source_role_refs == ["fc"]
    assert any("ghost-role" in w for w in warnings)


def test_validate_drops_unknown_outcome_ref() -> None:
    letter = _valid_letter(
        outcome_refs=[
            "Cut mobile load times from 10s to 2s",
            "Raised Series D of $1B",  # invented
        ]
    )
    cleaned, warnings = validate_cover_letter_refs(letter, _optimized_payload())
    assert cleaned.source_outcome_refs == ["Cut mobile load times from 10s to 2s"]
    assert any("Raised Series D" in w for w in warnings)


def test_validate_drops_unknown_skill_ref() -> None:
    letter = _valid_letter(skill_refs=["React", "COBOL"])
    cleaned, warnings = validate_cover_letter_refs(letter, _optimized_payload())
    assert cleaned.source_skill_refs == ["React"]
    assert any("COBOL" in w for w in warnings)


def test_validate_keeps_all_valid_refs() -> None:
    letter = _valid_letter()
    cleaned, warnings = validate_cover_letter_refs(letter, _optimized_payload())
    assert cleaned.source_role_refs == ["fc"]
    assert cleaned.source_outcome_refs == ["Cut mobile load times from 10s to 2s"]
    assert cleaned.source_skill_refs == ["React"]
    assert warnings == []


# ---------------------------------------------------------------------------
# tailor_cover_letter end-to-end
# ---------------------------------------------------------------------------


async def test_tailor_cover_letter_parses_and_returns_result() -> None:
    llm = MockLLMClient(scripted={DEFAULT_COVER_LETTER_PURPOSE: _valid_letter_json()})
    letter, warnings, result = await tailor_cover_letter(
        llm,
        optimized=_optimized_payload(),
        job_description="We want a senior FE",
        company_name="Acme",
        contact=_contact(),
    )
    assert isinstance(letter, TailoredCoverLetter)
    assert letter.recipient_company == "Acme Corp"
    assert len(letter.paragraphs) == 3
    assert result.cost_usd > 0
    assert warnings == []


async def test_tailor_cover_letter_cost_logs_under_default_purpose() -> None:
    llm = MockLLMClient(scripted={DEFAULT_COVER_LETTER_PURPOSE: _valid_letter_json()})
    await tailor_cover_letter(
        llm,
        optimized=_optimized_payload(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
    )
    assert llm.calls[0]["purpose"] == DEFAULT_COVER_LETTER_PURPOSE


async def test_tailor_cover_letter_enables_system_cache() -> None:
    llm = MockLLMClient(scripted={DEFAULT_COVER_LETTER_PURPOSE: _valid_letter_json()})
    await tailor_cover_letter(
        llm,
        optimized=_optimized_payload(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
    )
    assert llm.calls[0]["cache_system"] is True


async def test_tailor_cover_letter_drops_hallucinated_refs() -> None:
    llm = MockLLMClient(
        scripted={
            DEFAULT_COVER_LETTER_PURPOSE: _valid_letter_json(
                outcome_refs=["Cut mobile load times from 10s to 2s", "Invented win"],
                role_refs=["fc"],
                skill_refs=["React"],
            )
        }
    )
    letter, warnings, _ = await tailor_cover_letter(
        llm,
        optimized=_optimized_payload(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
    )
    assert letter.source_outcome_refs == ["Cut mobile load times from 10s to 2s"]
    assert any("Invented win" in w for w in warnings)


# ---------------------------------------------------------------------------
# render_cover_letter_docx
# ---------------------------------------------------------------------------


def _parse(doc_bytes: bytes) -> Any:
    return Document(io.BytesIO(doc_bytes))


def test_render_cover_letter_returns_valid_docx() -> None:
    letter = _valid_letter()
    doc = _parse(render_cover_letter_docx(letter))
    texts = [p.text for p in doc.paragraphs]
    assert "Daniel Joffe" in texts
    assert "Acme Corp" in texts
    assert any("Re: Senior FE" in t for t in texts)
    assert "Dear Acme Hiring Team," in texts
    assert any("FightCamp" in t for t in texts)
    assert "Sincerely," in texts


def test_render_cover_letter_omits_re_line_when_no_role() -> None:
    letter = _valid_letter().model_copy(update={"recipient_role": None})
    doc = _parse(render_cover_letter_docx(letter))
    texts = [p.text for p in doc.paragraphs]
    assert not any(t.startswith("Re:") for t in texts)


def test_render_cover_letter_has_no_tables() -> None:
    letter = _valid_letter()
    doc = _parse(render_cover_letter_docx(letter))
    assert len(doc.tables) == 0


def test_render_cover_letter_has_no_heading_1() -> None:
    letter = _valid_letter()
    doc = _parse(render_cover_letter_docx(letter))
    heading_1s = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert heading_1s == []


# ---------------------------------------------------------------------------
# lint_docx document_type awareness
# Heading-level rules moved to markdown_linter.py (see test_pandoc_smoke.py).
# ---------------------------------------------------------------------------


def test_lint_cover_letter_passes_byte_level_rules() -> None:
    letter = _valid_letter()
    result = lint_docx(render_cover_letter_docx(letter), document_type="cover_letter")
    assert result.ok is True
    assert result.violations == []


def test_lint_cover_letter_still_catches_tables() -> None:
    """Format rules apart from headings still apply to cover letters."""
    from docx import Document as NewDocument

    doc = NewDocument()
    doc.add_paragraph("body")
    doc.add_table(rows=1, cols=1)
    buf = io.BytesIO()
    doc.save(buf)
    result = lint_docx(buf.getvalue(), document_type="cover_letter")
    assert any(v.code == "no_tables" for v in result.errors)


# ---------------------------------------------------------------------------
# run_cover_letter_pipeline
# ---------------------------------------------------------------------------


async def test_pipeline_success_returns_record_and_uploads(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    supabase = _make_supabase_mock(insert_data=[_inserted_record_row()])
    monkeypatch.setattr(cost_log_mod, "record", MagicMock())

    llm = MockLLMClient(scripted={DEFAULT_COVER_LETTER_PURPOSE: _valid_letter_json()})
    result = await run_cover_letter_pipeline(
        supabase,
        llm,
        user_id="test-user",
        optimized=_optimized_doc(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
    )
    assert isinstance(result, CoverLetterPipelineSuccess)
    assert result.record.document_type == "cover_letter"
    supabase.storage.from_.assert_any_call("tailored-resumes")


async def test_pipeline_cost_logs_under_tailor_cover_letter(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    supabase = _make_supabase_mock(insert_data=[_inserted_record_row()])
    cost_record = MagicMock()
    monkeypatch.setattr(cost_log_mod, "record", cost_record)

    llm = MockLLMClient(scripted={DEFAULT_COVER_LETTER_PURPOSE: _valid_letter_json()})
    await run_cover_letter_pipeline(
        supabase,
        llm,
        user_id="test-user",
        optimized=_optimized_doc(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
    )
    assert cost_record.call_args.kwargs["purpose"] == DEFAULT_COVER_LETTER_PURPOSE
    assert (
        cost_record.call_args.kwargs["metadata"]["recipient_company"] == "Acme"
    )


async def test_pipeline_preferences_are_passed_through(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    supabase = _make_supabase_mock(insert_data=[_inserted_record_row()])
    monkeypatch.setattr(cost_log_mod, "record", MagicMock())

    seen: dict[str, str] = {}

    def responder(latest_user: str, _messages: object) -> str:
        seen["latest"] = latest_user
        return _valid_letter_json()

    llm = MockLLMClient(scripted={DEFAULT_COVER_LETTER_PURPOSE: responder})
    prefs = PreferencesPayload(
        rules=["lead with performance"],
        avoid=["em dashes"],
        tone_notes=[],
    )
    await run_cover_letter_pipeline(
        supabase,
        llm,
        user_id="test-user",
        optimized=_optimized_doc(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
        preferences=prefs,
    )
    assert "[Preferences]" in seen["latest"]
    assert "lead with performance" in seen["latest"]


async def test_pipeline_lint_failure_does_not_persist(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    supabase = _make_supabase_mock(insert_data=[])
    monkeypatch.setattr(cost_log_mod, "record", MagicMock())

    def fake_lint(_b: bytes, *, document_type: str = "resume") -> LintResult:
        return LintResult(
            ok=False,
            violations=[
                LintViolation(
                    code="no_tables",
                    message="simulated lint failure",
                    severity="error",
                )
            ],
        )

    monkeypatch.setattr(pipeline_module, "lint_docx", fake_lint)

    llm = MockLLMClient(scripted={DEFAULT_COVER_LETTER_PURPOSE: _valid_letter_json()})
    result = await run_cover_letter_pipeline(
        supabase,
        llm,
        user_id="test-user",
        optimized=_optimized_doc(),
        job_description="jd",
        company_name="Acme",
        contact=_contact(),
    )
    assert isinstance(result, CoverLetterPipelineLintFailure)
    supabase.table.return_value.insert.assert_not_called()


# ---------------------------------------------------------------------------
# Regression: resume pipeline still works after P5 changes
# ---------------------------------------------------------------------------


def test_resume_pipeline_still_lints_clean() -> None:
    """After P5 added document_type, render_docx + lint_docx default path
    should stay green on a clean resume."""
    from app.models.tailor import (
        TailoredBullet,
        TailoredEducation,
        TailoredResume,
        TailoredRole,
    )

    resume = TailoredResume(
        summary="Senior FE with shipped work.",
        contact=_contact(),
        experience=[
            TailoredRole(
                company="FightCamp",
                title="Senior FE",
                start="2021-11",
                end="2024-04",
                bullets=[
                    TailoredBullet(
                        text="Cut mobile load times from 10s to 2s.",
                        source_outcome_ref="x",
                    )
                ],
                source_role_ref="fc",
            )
        ],
        skills=["React"],
        education=[TailoredEducation(school="UCLA")],
    )
    result = lint_docx(render_docx(resume))
    assert result.ok is True
