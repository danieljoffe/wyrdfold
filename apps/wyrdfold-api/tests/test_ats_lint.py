"""ATS linter tests. Happy path via the real renderer; failure modes
constructed directly with python-docx to simulate regressions.
"""

from __future__ import annotations

import io
import zipfile
from typing import Any

from docx import Document

from app.models.tailor import (
    ContactInfo,
    TailoredBullet,
    TailoredEducation,
    TailoredResume,
    TailoredRole,
)
from app.services.ats_lint import lint_docx
from app.services.docx.renderer import render_docx


def _good_resume() -> TailoredResume:
    return TailoredResume(
        summary="Senior FE with a decade of shipped work.",
        contact=ContactInfo(name="Daniel Joffe", email="daniel@example.com"),
        experience=[
            TailoredRole(
                company="FightCamp",
                title="Senior Frontend Engineer",
                start="2021-11",
                end="2024-04",
                bullets=[
                    TailoredBullet(
                        text="Cut mobile load times from 10s to 2s.",
                        source_outcome_ref="x",
                    )
                ],
                source_role_ref="fc",
            )
        ],
        skills=["React", "TypeScript"],
        education=[TailoredEducation(school="UCLA")],
    )


def _doc_to_bytes(doc: Any) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _minimum_valid_doc() -> Any:
    """A baseline doc with just the required Experience heading so the
    experience_heading rule passes; other rules stay clean.
    """
    doc = Document()
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Role · Company · 2020 - present")
    return doc


def _inject_zip_entry(data: bytes, path: str, contents: bytes) -> bytes:
    """Append a new entry into a .docx zip. Used to simulate
    regressions (e.g. stray word/media entries) without forcing
    python-docx to accept synthetic images.
    """
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                zout.writestr(item, zin.read(item.filename))
            zout.writestr(path, contents)
    return out.getvalue()


# ---- Happy path ----------------------------------------------------------


def test_good_resume_lints_clean() -> None:
    result = lint_docx(render_docx(_good_resume()))
    assert result.ok is True
    assert result.errors == []
    assert result.warnings == []


# ---- valid_docx -----------------------------------------------------------


def test_invalid_bytes_fails_valid_docx() -> None:
    result = lint_docx(b"not a docx")
    assert result.ok is False
    codes = {v.code for v in result.violations}
    assert "valid_docx" in codes


def test_empty_bytes_fails_valid_docx() -> None:
    result = lint_docx(b"")
    assert result.ok is False


# ---- no_tables -----------------------------------------------------------


def test_table_triggers_no_tables_error() -> None:
    doc = _minimum_valid_doc()
    doc.add_table(rows=1, cols=2)
    result = lint_docx(_doc_to_bytes(doc))
    assert result.ok is False
    assert any(v.code == "no_tables" for v in result.errors)


# ---- no_inline_images ----------------------------------------------------


def test_inline_image_triggers_no_inline_images_error() -> None:
    """Simulate a regression by injecting a word/media/ zip entry
    directly. This is what `python-docx.add_picture` would produce,
    without depending on the library accepting synthetic image bytes.
    """
    base = _doc_to_bytes(_minimum_valid_doc())
    with_image = _inject_zip_entry(base, "word/media/image1.png", b"fake-png")
    result = lint_docx(with_image)
    assert result.ok is False
    assert any(v.code == "no_inline_images" for v in result.errors)


# Heading-related rules were moved to markdown_linter.py — see
# tests/test_pandoc_smoke.py for their coverage. The docx-byte linter
# is now strictly the byte-level safety net (tables, images, frames,
# shapes, page_count).


# ---- page_count ----------------------------------------------------------


def test_short_doc_has_no_page_count_violation() -> None:
    result = lint_docx(render_docx(_good_resume()))
    assert not any(v.code == "page_count" for v in result.violations)


def test_medium_doc_warns_on_page_count() -> None:
    doc = _minimum_valid_doc()
    for i in range(85):
        doc.add_paragraph(f"filler paragraph {i}")
    result = lint_docx(_doc_to_bytes(doc))
    codes = {v.code for v in result.warnings}
    assert "page_count" in codes
    assert result.ok is True  # warn, not error


def test_large_doc_errors_on_page_count() -> None:
    doc = _minimum_valid_doc()
    for i in range(130):
        doc.add_paragraph(f"filler paragraph {i}")
    result = lint_docx(_doc_to_bytes(doc))
    codes = {v.code for v in result.errors}
    assert "page_count" in codes
    assert result.ok is False


# ---- LintResult shape ----------------------------------------------------


def test_result_errors_and_warnings_partition_violations() -> None:
    doc = _minimum_valid_doc()
    doc.add_table(rows=1, cols=1)  # error: no_tables
    for i in range(85):  # warning: page_count (>80 lines)
        doc.add_paragraph(f"filler paragraph {i}")
    result = lint_docx(_doc_to_bytes(doc))
    assert len(result.errors) == 1
    assert len(result.warnings) == 1
    assert result.errors[0].severity == "error"
    assert result.warnings[0].severity == "warning"
    assert len(result.errors) + len(result.warnings) == len(result.violations)


def test_ok_flag_mirrors_presence_of_errors() -> None:
    clean = lint_docx(render_docx(_good_resume()))
    assert clean.ok is True

    broken = _minimum_valid_doc()
    broken.add_table(rows=1, cols=1)
    assert lint_docx(_doc_to_bytes(broken)).ok is False


def test_violation_codes_are_stable_machine_ids() -> None:
    doc = _minimum_valid_doc()
    doc.add_table(rows=1, cols=1)
    result = lint_docx(_doc_to_bytes(doc))
    # Codes are snake_case, lowercase, no spaces.
    for v in result.violations:
        assert v.code == v.code.lower()
        assert " " not in v.code
