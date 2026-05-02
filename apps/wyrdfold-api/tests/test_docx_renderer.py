"""ATS-friendly .docx renderer tests.

Inspects the generated document structurally rather than comparing bytes
— python-docx writes some timestamps into core.xml which would break
byte-equal golden tests. The content we care about lives in document.xml
and is stable.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document

from app.models.tailor import (
    ContactInfo,
    TailoredBullet,
    TailoredEducation,
    TailoredResume,
    TailoredRole,
)
from app.services.docx.renderer import (
    SECTION_EDUCATION,
    SECTION_EXPERIENCE,
    SECTION_SKILLS,
    SECTION_SUMMARY,
    format_date,
    render_docx,
)

# ---- format_date ----------------------------------------------------------


def test_format_date_converts_month() -> None:
    assert format_date("2021-11") == "Nov 2021"


def test_format_date_none_returns_present() -> None:
    assert format_date(None) == "Present"


def test_format_date_invalid_echoes_input() -> None:
    assert format_date("whenever") == "whenever"


def test_format_date_out_of_range_echoes_input() -> None:
    assert format_date("2021-13") == "2021-13"


# ---- Fixture resume -------------------------------------------------------


def _resume() -> TailoredResume:
    return TailoredResume(
        summary="Senior FE with a decade of shipped work.",
        contact=ContactInfo(
            name="Daniel Joffe",
            email="daniel@example.com",
            location="Los Angeles, CA",
            linkedin="linkedin.com/in/daniel",
        ),
        experience=[
            TailoredRole(
                company="FightCamp",
                title="Senior Frontend Engineer",
                start="2021-11",
                end="2024-04",
                bullets=[
                    TailoredBullet(
                        text="Cut mobile load times from 10s to 2s.",
                        source_outcome_ref="Cut mobile load times",
                    ),
                    TailoredBullet(
                        text="Led the PDP rebuild.",
                        source_outcome_ref="PDP",
                    ),
                ],
                source_role_ref="fc",
            )
        ],
        skills=["React", "TypeScript", "Next.js"],
        education=[
            TailoredEducation(
                school="UCLA", degree="BA, Design Media Arts", dates="2010 - 2014"
            )
        ],
    )


def _parse(doc_bytes: bytes) -> Any:
    return Document(io.BytesIO(doc_bytes))


# ---- Output validity ------------------------------------------------------


def test_render_returns_non_empty_bytes() -> None:
    assert len(render_docx(_resume())) > 0


def test_render_output_opens_as_valid_docx() -> None:
    # Would raise if the bytes aren't a valid .docx archive.
    _parse(render_docx(_resume()))


# ---- Content fidelity -----------------------------------------------------


def test_name_is_title() -> None:
    doc = _parse(render_docx(_resume()))
    paragraphs = [p.text for p in doc.paragraphs]
    assert paragraphs[0] == "Daniel Joffe"


def test_contact_line_includes_location_and_email() -> None:
    doc = _parse(render_docx(_resume()))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Los Angeles, CA" in all_text
    assert "daniel@example.com" in all_text


def test_summary_section_present() -> None:
    doc = _parse(render_docx(_resume()))
    texts = [p.text for p in doc.paragraphs]
    assert SECTION_SUMMARY in texts
    assert any("Senior FE" in t for t in texts)


def test_experience_section_present() -> None:
    doc = _parse(render_docx(_resume()))
    texts = [p.text for p in doc.paragraphs]
    assert SECTION_EXPERIENCE in texts
    assert any("FightCamp" in t for t in texts)
    assert any("Senior Frontend Engineer" in t for t in texts)


def test_role_dates_formatted_human() -> None:
    doc = _parse(render_docx(_resume()))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Nov 2021" in all_text
    assert "Apr 2024" in all_text


def test_bullets_rendered_as_list_bullet_style() -> None:
    doc = _parse(render_docx(_resume()))
    bullet_texts = [
        p.text for p in doc.paragraphs if p.style.name == "List Bullet"
    ]
    assert "Cut mobile load times from 10s to 2s." in bullet_texts
    assert "Led the PDP rebuild." in bullet_texts


def test_skills_section_flat_comma_list() -> None:
    doc = _parse(render_docx(_resume()))
    texts = [p.text for p in doc.paragraphs]
    assert SECTION_SKILLS in texts
    assert "React, TypeScript, Next.js" in texts


def test_education_section_present() -> None:
    doc = _parse(render_docx(_resume()))
    texts = [p.text for p in doc.paragraphs]
    assert SECTION_EDUCATION in texts
    assert any("UCLA" in t for t in texts)
    assert any("BA, Design Media Arts" in t for t in texts)


# ---- ATS constraints ------------------------------------------------------


def test_no_tables_used() -> None:
    doc = _parse(render_docx(_resume()))
    assert len(doc.tables) == 0


def test_no_inline_images_used() -> None:
    doc = _parse(render_docx(_resume()))
    blob = io.BytesIO()
    doc.save(blob)
    raw = blob.getvalue()
    # .docx inline images land as word/media/ entries in the zip.
    assert b"word/media" not in raw


def test_no_text_frames_used() -> None:
    doc = _parse(render_docx(_resume()))
    xml = doc.element.xml
    # wps:txbx (word processing shapes text box) would break ATS parsing.
    assert "txbx" not in xml


# ---- Conditional sections -------------------------------------------------


def test_empty_experience_omits_experience_section() -> None:
    resume = TailoredResume(
        summary="summary text",
        contact=ContactInfo(name="x"),
        experience=[],
        skills=["React"],
    )
    doc = _parse(render_docx(resume))
    texts = [p.text for p in doc.paragraphs]
    assert SECTION_EXPERIENCE not in texts


def test_empty_skills_omits_skills_section() -> None:
    resume = TailoredResume(
        summary="summary text",
        contact=ContactInfo(name="x"),
        experience=[],
        skills=[],
    )
    doc = _parse(render_docx(resume))
    texts = [p.text for p in doc.paragraphs]
    assert SECTION_SKILLS not in texts


def test_empty_education_omits_education_section() -> None:
    resume = TailoredResume(
        summary="summary text",
        contact=ContactInfo(name="x"),
        experience=[],
        skills=["React"],
        education=[],
    )
    doc = _parse(render_docx(resume))
    texts = [p.text for p in doc.paragraphs]
    assert SECTION_EDUCATION not in texts


def test_current_role_shows_present() -> None:
    resume = _resume()
    resume.experience[0] = resume.experience[0].model_copy(update={"end": None})
    doc = _parse(render_docx(resume))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Present" in all_text


# ---- Minimal contact ------------------------------------------------------


def test_render_works_with_name_only_contact() -> None:
    resume = TailoredResume(
        summary="summary",
        contact=ContactInfo(name="Daniel Joffe"),
        experience=[],
        skills=[],
    )
    doc = _parse(render_docx(resume))
    paragraphs = [p.text for p in doc.paragraphs]
    assert paragraphs[0] == "Daniel Joffe"
