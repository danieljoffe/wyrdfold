"""Deterministic ATS format linter over rendered `.docx` bytes.

Markdown is the source of truth for the document structure; the
markdown linter (markdown_linter.py) owns the heading/section rules.
This module is the byte-level safety net — it catches issues the
markdown text can't see (text frames, image entries, shape XML,
malformed archive).

Rules (Greenhouse-floor):
- valid_docx (error)          — opens as a valid zipped OOXML package
- no_tables (error)           — `doc.tables` empty
- no_inline_images (error)    — no `word/media/*` entries in the zip
- no_text_frames (error)      — no `txbx` in document XML
- no_shapes (error)           — no `w:drawing` / `w:pict` in document XML
- page_count (warning/error)  — ~paragraph count heuristic: > 80 warn, > 120 error
"""

from __future__ import annotations

import io
import zipfile
from typing import Literal

from docx import Document

from app.models.ats_lint import LintResult, LintViolation

_PAGE_COUNT_WARN_AT = 80
_PAGE_COUNT_ERROR_AT = 120


def _parse_or_none(data: bytes) -> object | None:
    try:
        return Document(io.BytesIO(data))
    except Exception:
        return None


def _raw_xml_or_empty(doc: object) -> str:
    try:
        return str(doc.element.xml)  # type: ignore[attr-defined]
    except Exception:
        return ""


def _zip_entries(data: bytes) -> list[str]:
    try:
        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            return zf.namelist()
    except zipfile.BadZipFile:
        return []


def _non_empty_paragraph_count(doc: object) -> int:
    paragraphs = doc.paragraphs  # type: ignore[attr-defined]
    return len([p for p in paragraphs if p.text.strip()])


def lint_docx(
    data: bytes,
    *,
    document_type: Literal["resume", "cover_letter"] = "resume",
) -> LintResult:
    violations: list[LintViolation] = []

    entries = _zip_entries(data)
    doc = _parse_or_none(data)
    if doc is None or not entries:
        violations.append(
            LintViolation(
                code="valid_docx",
                message="Bytes are not a valid .docx archive.",
                severity="error",
            )
        )
        return LintResult(ok=False, violations=violations)

    # no_inline_images
    if any(name.startswith("word/media/") for name in entries):
        violations.append(
            LintViolation(
                code="no_inline_images",
                message=(
                    "Document contains inline images (word/media/ entries). "
                    "ATS parsers routinely drop image content."
                ),
                severity="error",
            )
        )

    # no_tables
    if len(doc.tables) > 0:  # type: ignore[attr-defined]
        violations.append(
            LintViolation(
                code="no_tables",
                message=(
                    "Document contains tables. Most ATS parsers read tables "
                    "inconsistently; use single-column paragraph content."
                ),
                severity="error",
            )
        )

    xml = _raw_xml_or_empty(doc)
    # no_text_frames
    if "txbx" in xml:
        violations.append(
            LintViolation(
                code="no_text_frames",
                message=(
                    "Document contains text frames (txbx). Text inside frames "
                    "often does not survive ATS parsing."
                ),
                severity="error",
            )
        )
    # no_shapes
    if "w:drawing" in xml or "w:pict" in xml:
        violations.append(
            LintViolation(
                code="no_shapes",
                message=(
                    "Document contains drawings or picture shapes. These are "
                    "unreadable to ATS parsers."
                ),
                severity="error",
            )
        )

    # page_count
    paragraph_count = _non_empty_paragraph_count(doc)
    if paragraph_count > _PAGE_COUNT_ERROR_AT:
        violations.append(
            LintViolation(
                code="page_count",
                message=(
                    f"Document has {paragraph_count} non-empty paragraphs; "
                    f"likely exceeds 3 pages. Tighten content."
                ),
                severity="error",
            )
        )
    elif paragraph_count > _PAGE_COUNT_WARN_AT:
        violations.append(
            LintViolation(
                code="page_count",
                message=(
                    f"Document has {paragraph_count} non-empty paragraphs; "
                    f"likely exceeds 2 pages."
                ),
                severity="warning",
            )
        )

    ok = not any(v.severity == "error" for v in violations)
    return LintResult(ok=ok, violations=violations)
